# merni listi
# by
# Daniel Nikolov

import os
import time
from functools import reduce
import creopyson
from pathlib import Path
import getpass
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

#v1.0
#merni listi da pravi dokument i so klikanje da dobivame koti

#v1.1
# popraven bug da ne otvara crtez i
# da ne pagja programata na poslozeni ctrezi

#v1.2
# popraven bug da gi cita i superscript kotite
# promeneto filename na merna lista da bide isto kako drawing name

#V1.3
# popraveno da dodava ista kota (ako e kliknata poveke pati)
# center na sekoja kelija

#v1.4
# promenet templejt
# dodadena redica za vnesuvanje naziv i kolicina

################ funkcii

var_brojac_new = 1
var_brojac = var_brojac_new
var_brojac_new = str(var_brojac_new)

def save_document():
    ############################################### making document!
    document = Document(template_path_ml)
    section = document.sections[0]
    header = section.header
    footer = section.footer
    p = document.paragraphs

    for section in document.sections:
        footer.paragraphs[1].text = footer.paragraphs[1].text.replace(var_poz_old, var_poz_new)
        footer.paragraphs[1].text = footer.paragraphs[1].text.replace(var_naziv_old, var_naziv_new)
        header.paragraphs[2].text = header.paragraphs[2].text.replace(var_orodje_old, var_orodje_new)
        header.paragraphs[3].text = header.paragraphs[3].text.replace(var_koda_old, var_koda_new)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if var_naziv_old in paragraph.text:
                        paragraph.text = paragraph.text.replace(var_naziv_old, var_naziv_new)
                    if var_brojac_old in paragraph.text:
                        paragraph.text = paragraph.text.replace(var_brojac_old, var_brojac_new)
                    if var_kolicina_old in paragraph.text:
                        paragraph.text = paragraph.text.replace(var_kolicina_old, var_kolicina_new)

    j = 1
    k = 2
    for i in list_sort:
        name = (i["name"])
        value = (i["value"])
        value = round(value, 2)
        try:
            tolerance_type = (i["tolerance_type"])
        except Exception as e:
            tolerance_type = (i["dim_type"])

        text = (i['text'])

        text_tmp = str(text)
        # print(text_tmp)
        if text_tmp == r"['{0:@D}\n']":
            text_prefix = ""
        elif text_tmp == r"['{0:\x01n\x02}{1:@D}\n']":
            text_prefix = "Ø"
        elif text_tmp == r"['{0:M}{1:@D}\n']":
            text_prefix = "M"
        elif text_tmp == r"['{0:R}{1:@D}\n']":
            text_prefix = "R"
        else:
            text_prefix = ""
        if (i["dim_type"]) == "angular":
            text_suffix = "°"
        else:
            text_suffix = ""
        if tolerance_type == 'isodin':
            tolerance = (i['tol_table_name'])  ## H, g, F, e...
            if tolerance != 'NONE':
                tolerance_value = (i["tol_table_column"])  ## 7, 6...
            else:
                tolerance = ""
                tolerance_value = ""
            tolerancija = f"{tolerance}{tolerance_value}"
        elif tolerance_type == 'symmetric':
            tolerance = "±"  ## ±
            tolerance_value = (i["tol_symmetric_value"])  ## 7, 6...
            tolerance_value = round(tolerance_value, 3)
            tolerancija = f"{tolerance}{tolerance_value}"
        elif tolerance_type == 'sym_superscript':
            tolerance = "±"  ## ±
            tolerance_value = (i["tol_symmetric_value"])  ## 7, 6...
            tolerance_value = round(tolerance_value, 3)
            tolerancija = f"{tolerance}{tolerance_value}"
        elif tolerance_type == 'plus_minus':
            tol_minus = (i['tol_minus'])  ## vrednosti
            tol_plus = (i['tol_plus'])  ## vrednosti
            if tol_minus <= 0:
                tolerance_minus = "+"
                tol_minus = abs(tol_minus)
                tol_minus = round(tol_minus, 3)
            else:
                tolerance_minus = "-"
                tol_minus = round(tol_minus, 3)
            if tol_plus >= 0:
                tolerance_plus = "+"
                tol_plus = round(tol_plus, 3)
            else:
                tol_plus = abs(tol_plus)
                tol_plus = round(tol_plus, 3)
                tolerance_plus = "-"
            tolerancija = f"""{tolerance_plus}{tol_plus}
{tolerance_minus}{tol_minus}"""
        else:
            tolerance = ""
            tolerance_value = ""
            tolerancija = f"{tolerance}{tolerance_value}"
        tolerancija = str(tolerancija)
        value = str(value)
        value = f"{text_prefix}{value}{text_suffix}"
        # print(name)
        # print(value)
        # print(tolerancija)
        # print(j)
        document.tables[0].add_row()  # ADD ROW HERE
        document.tables[0].cell(k, 0).text = str(j)
        document.tables[0].cell(k, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.tables[0].cell(k, 1).text = value
        document.tables[0].cell(k, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.tables[0].cell(k, 2).text = tolerancija
        document.tables[0].cell(k, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        j = j + 1
        k = k + 1
    #delete last row
    last_row = k

    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    table = document.tables[0]
    row = table.rows[last_row]
    remove_row(table, row)

    font_1 = document.styles['Normal'].font
    font_1.name = 'Arial'
    font_1.bold = True
    font_2 = document.styles["Header"].font
    font_2.name = 'Arial'
    font_2.bold = True
    font_3 = document.styles["Footer"].font
    font_3.name = 'Arial'
    font_3.bold = True

    # font.size = docx.shared.Pt(12)
    document.save(merni_listi_pateka + dokument)
    print("document saved at: ", merni_listi_pateka)
    return
###############################################

c = creopyson.Client()
user = getpass.getuser()

# read me
print("""
This program is for creating document "Measurement lists"
IMPORTANT NOTE:
IF MEASUREMENT LIST EXIST, IT WILL BE OVERWRITTEN!
1. You need to have opened Creo.
if you don't have opened yet, you can do it now.
2. Open the drawing from which you want to create measurement list  
3. Select all dimensions for measuring
4. When you are done, simply click middle mouse button
5. if you want to add new measurement list FIRST OPEN NEW DRAWING then press "y" 
""")


CREOSON_PATH = Path(f"C:\\Users\\{user}\\Desktop\\creoson")
image_lth_path = Path(f"C:\\Users\\{user}\\Documents\\lth_logo\\lth_logo.png")
template_path_ml = Path(f"C:\\Users\\{user}\\Documents\\merni_listi_template\\ML_template.docx")
os.chdir(CREOSON_PATH)
run_path = Path(CREOSON_PATH) / "creoson_run.bat"

os.startfile(run_path.resolve().as_posix())
# variables for document
var_orodje_old = "ORODJE"
var_koda_old = "KODA"
var_poz_old = "POZ"
var_naziv_old = "NAZIV"
var_brojac_old = "BROJAC"
var_kolicina_old = "KOLICINA"

c.connect()
while c.is_creo_running() == False:
    print("""
--------------------
Creo is not running!
Please start Creo!
--------------------""")
    time.sleep(10)
else:
    print("Creo is running!")


current_directory = c.creo_pwd()
merni_listi = "Merni listi"
pateka = Path(f"{current_directory}{merni_listi}")  #  create folder ML in WD
if not Path.is_dir(pateka):
    Path.mkdir(pateka)
    print("Merni listi folder was made!")
else:
    print("Merni listi folder exists!")
    pateka.rename(Path(current_directory, merni_listi))
new_file_list = ("", " ", "yes", "y", "Y", "YES")
new_file = ""
#check if is drw
while new_file in new_file_list:
    check_is_drw = True
    while check_is_drw == True:
        try:
            file_get_active = c.file_get_active()
            file_get_active = file_get_active["file"]
            is_drw = file_get_active[-4:]
            check_is_drw = False
        except Exception as e:
            check_is_drw = True
            print("Error occurred: ", e)
            print("open drawing!")
            time.sleep(10)

    while is_drw != ".drw":
        print("Open drawing")
        time.sleep(10)
        file_get_active = c.file_get_active()
        file_get_active = file_get_active["file"]
        is_drw = file_get_active[-4:]

    file_get_active_new = file_get_active

    print("Active drawing:", file_get_active_new)
    model_file = c.drawing_get_cur_model()
    parametar_pozcija_raw = c.parameter_list("poz", file_=model_file)
    parametar_orodje_raw = c.parameter_list("orodje", file_=model_file)
    parametar_koda_raw = c.parameter_list("koda", file_=model_file)
    parametar_naziv_raw = c.parameter_list("naziv", file_=model_file)
    parametar_st_kosov_raw = c.parameter_list("st_kosov", file_=model_file)

    var_poz_new = reduce(lambda a, b: dict(a, **b), parametar_pozcija_raw)
    var_poz_new = (var_poz_new["value"])
    if var_poz_new == 0:
        var_poz_new = input("enter 'POZ' parameter: ")
    var_poz_new = str(var_poz_new)
    var_orodje_new = reduce(lambda a, b: dict(a, **b), parametar_orodje_raw)
    var_orodje_new = (var_orodje_new["value"])
    var_koda_new = reduce(lambda a, b: dict(a, **b), parametar_koda_raw)
    var_koda_new = (var_koda_new["value"])
    var_naziv_new = reduce(lambda a, b: dict(a, **b), parametar_naziv_raw)
    var_naziv_new = (var_naziv_new["value"])
    var_kolicina_new = reduce(lambda a, b: dict(a, **b), parametar_st_kosov_raw)
    var_kolicina_new = (var_kolicina_new["value"])
    var_kolicina = var_kolicina_new
    if var_kolicina_new == 0:
        var_kolicina_new = int(input("Insert value for 'quantity': "))
    var_kolicina_new = str(var_kolicina_new)

    dokument = ".docx"
    ml_name = file_get_active_new[:-4]
    merni_listi_pateka = f"{pateka}\\{var_poz_new}_{ml_name}"

    list_dimensions = c.dimension_list_detail()
    print("All Dimensions on drawing:", len(list_dimensions))
    #print(list_dimensions)
    list_name_dim = []
    list_value_dim = []

    break_point = True
    print("Select dimensions!")
    while break_point == True:
        try:
            dimension_select = c.dimension_user_select()
            #print(dimension_select)
            dimension_raw = reduce(lambda a, b: dict(a, **b), dimension_select)
            dimension_name = dimension_raw["name"]
            dimension_value = str(round(dimension_raw["value"], 2))
            list_name_dim.append(dimension_name)
            list_value_dim.append(dimension_value)

            print("Selected dimension: " + dimension_value)

        except Exception as e:
            #print(e)

            break_point = False

    print("number of selected dimensions:", len(list_value_dim))
    #print("list of chosen dimensions:", list_value_dim)
    list_filter = []
    for x in list_name_dim:
        for y in list_dimensions:
            if y["name"] == x:
                list_filter.append(y)
    #print(list_filter)
#funkcionira samo bez dupli koti
    #list_filter = [x for x in list_dimensions if x["name"] in list_name_dim]
    #number_columns = len(list_filter)
    #list_sort = sorted(list_filter, key=lambda item: list_name_dim.index(item['name']))
##############

    list_sort = list_filter

    #print(list_sort)
    # print(list_filter)
    # print(list_name_dim)
    # print(list_sort)
    save_document()
    print("")
    print("For new measurement list first open new drawing then press 'Y' for stopping the program press 'N'!")
    new_file = input("Enter 'Y' or 'N': ")

print("done")
time.sleep(20)
